import pytest
import os
from backend.services.importer.docx_importer import DOCXImporter
import docx

@pytest.fixture
def sample_docx(tmp_path):
    """Creates a temporary sample DOCX file for testing."""
    file_path = tmp_path / "sample_resume.docx"
    doc = docx.Document()
    
    # Contact
    doc.add_paragraph("John Doe")
    doc.add_paragraph("johndoe@example.com | +1234567890")
    
    # Summary
    doc.add_paragraph("Summary")
    doc.add_paragraph("Experienced software engineer with a passion for building scalable backends.")
    
    # Experience
    doc.add_paragraph("Experience")
    doc.add_paragraph("Tech Corp - Senior Engineer (2020 - Present)")
    doc.add_paragraph("Built microservices using Python and Go.")
    
    doc.save(file_path)
    return str(file_path)

def test_docx_importer_extraction(sample_docx):
    importer = DOCXImporter()
    profile, confidence = importer.import_resume(sample_docx)
    
    # Test Contact Extraction
    assert profile.contact.full_name == "John Doe"
    assert profile.contact.email == "johndoe@example.com"
    
    # Test Summary
    assert "Experienced software engineer" in profile.profile.summary
    
    # Test Experience
    assert len(profile.experience) == 1
    assert "Tech Corp" in profile.experience[0].responsibilities[0]
    
    # Test Confidence Engine
    assert confidence["contact"] == 0.80
    assert confidence["profile"] > 0.50
    assert confidence["experience"] > 0.50
    assert confidence["projects"] == 0.0  # Not in doc
